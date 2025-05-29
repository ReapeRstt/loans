import tkinter
import random
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime, timedelta
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
import mysql.connector

class Add_en():
    def __init__(self):
        def add_to_db():
            name = name_entry.get()
            address = address_entry.get()
            phone = phone_entry.get()
            inn = inn_entry.get()

            # Установка соединения с базой данных MySQL
            conn = mysql.connector.connect(
                host='localhost',
                user='root',
                password='22345267',
                database='loan_db'
            )

            # Создание курсора для выполнения SQL-запросов
            cursor = conn.cursor()

            # Выполнение SQL-запроса для добавления данных в таблицу legal_entities
            cursor.execute(
                "INSERT INTO legal_entities (en_name, address, phone_number, inn) "
                "VALUES (%s, %s, %s, %s)",
                (name, address, phone, inn))

            # Применение изменений
            conn.commit()

            def call_update(self):
                self.update_object.update_treeview()

            # Закрытие соединения с базой данных
            conn.close()

            messagebox.showinfo("Добавление.", "Запись занесена в базу данные.")

        def format_phone(event=None):
            # Форматирует номер по маске ##-##-##
            current = phone_entry.get().replace("-", "")
            formatted = ""

            if len(current) > 0:
                formatted = current[:2]  # Первые 2 цифры
            if len(current) > 2:
                formatted += "-" + current[2:4]  # Добавляем - и следующие 2 цифры
            if len(current) > 4:
                formatted += "-" + current[4:6]  # Добавляем - и последние 2 цифры

            # Обновляем поле ввода
            phone_entry.delete(0, "end")
            phone_entry.insert(0, formatted)

        def validate_inn_input(new_value):
            "Проверка ввода ИНН (только цифры, максимум 12 символов)"
            return new_value.isdigit() and len(new_value) <= 12 or new_value == ""

        entitie_add = Tk()
        entitie_add.title("Добавление юр. лица.")
        entitie_add['bg'] = '#d7cecc'
        entitie_add.geometry('350x300')
        entitie_add.resizable(width=False, height=False)

        name_label = Label(entitie_add, text="Название", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        name_label.pack()
        name_entry = Entry(entitie_add, bg='#f0f6f6')
        name_entry.pack()

        address_label = Label(entitie_add, text="Адрес", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        address_label.pack()
        address_entry = Entry(entitie_add, bg='#f0f6f6')
        address_entry.pack()

        phone_label = Label(entitie_add, text="Номер телефона", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        phone_label.pack()
        phone_entry = Entry(entitie_add, bg='#f0f6f6')
        phone_entry.bind("<KeyRelease>", format_phone)
        phone_entry.pack()

        inn_label = Label(entitie_add, text="ИНН", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        inn_label.pack()
        inn_entry = Entry(entitie_add, bg='#f0f6f6', validate="key")
        inn_entry['validatecommand'] = (entitie_add.register(validate_inn_input), '%P')
        inn_entry.pack()

        add_button = Button(entitie_add, text="Добавить в базу данных", command=add_to_db, bg='#f0f6f6')
        add_button.pack(padx=10, pady=8)


class Add_lo():
    def __init__(self, entity_id=""):
        # Создаем окно
        self.loan_add = Tk()
        self.loan_add.title("Добавление кредита.")
        self.loan_add['bg'] = '#d7cecc'
        self.loan_add.geometry('350x300')
        self.loan_add.resizable(width=False, height=False)

        # Создаем элементы интерфейса
        en_id_label = Label(self.loan_add, text="ID лица", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        en_id_label.pack()
        self.en_id_entry = Entry(self.loan_add, bg='#f0f6f6')
        self.en_id_entry.pack()

        # Если передан ID юрлица, подставляем его в поле
        if entity_id:
            self.en_id_entry.insert(0, entity_id)

        amount_label = Label(self.loan_add, text="Сумма", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        amount_label.pack()
        self.amount_entry = Entry(self.loan_add, bg='#f0f6f6')
        self.amount_entry.pack()

        percent_label = Label(self.loan_add, text="Процент", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        percent_label.pack()
        self.percent_entry = Entry(self.loan_add, bg='#f0f6f6')
        self.percent_entry.pack()

        term_label = Label(self.loan_add, text="Срок", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
        term_label.pack()
        self.term_entry = Entry(self.loan_add, bg='#f0f6f6')
        self.term_entry.pack()

        add_button = Button(self.loan_add, text="Добавить в базу данных", command=self.add_to_db, bg='#f0f6f6')
        add_button.pack(padx=10, pady=8)

    def calculate_payments(self, amount, percent, term):
        """Расчёт финансовых условий для юр.лиц с проверкой типов"""
        try:
            # Преобразуем введенные значения в числа
            amount = float(amount.replace(' ', '').replace(',', '.'))
            percent = float(percent.replace(' ', '').replace(',', '.'))
            term = int(term)

            # Проверка на валидность значений
            if amount <= 0 or percent <= 0 or term <= 0:
                raise ValueError("Значения должны быть положительными")

            # Профессиональная формула аннуитетного платежа
            monthly_rate = percent / 100 / 12
            annuity_coeff = (monthly_rate * (1 + monthly_rate) ** term) / ((1 + monthly_rate) ** term - 1)
            monthly_payment = round(amount * annuity_coeff, 2)

            # Начальный взнос 15-30% для юр.лиц
            initial_payment_percent = 20 + random.randint(-5, 10)  # Вариация 15-30%
            initial_payment = round(amount * initial_payment_percent / 100, 2)

            # Общая переплата
            total_payment = round(monthly_payment * term, 2)
            total_overpayment = round(total_payment - amount, 2)

            return {
                'initial_payment': initial_payment,
                'initial_payment_percent': initial_payment_percent,
                'monthly_payment': monthly_payment,
                'total_payment': total_payment,
                'total_overpayment': total_overpayment
            }
        except ValueError as e:
            messagebox.showerror("Ошибка данных", f"Некорректные входные данные: {str(e)}")
            return None
        except Exception as e:
            messagebox.showerror("Ошибка расчета", f"Ошибка при расчете платежей: {str(e)}")
            return None

    def generate_credit_contract(self, loan_id, en_id, client_data, amount, percent, term):
        """Генерация кредитного договора с профессиональными финансовыми расчетами"""
        try:
            # Проверка и преобразование входных данных
            try:
                amount_clean = float(str(amount).replace(' ', '').replace(',', '.'))
                percent_clean = float(str(percent).replace(' ', '').replace(',', '.'))
                term_clean = int(str(term).replace(' ', ''))

                if amount_clean <= 0 or percent_clean <= 0 or term_clean <= 0:
                    raise ValueError("Все значения должны быть положительными числами")
            except ValueError as e:
                messagebox.showerror("Ошибка данных", f"Некорректные данные: {str(e)}")
                return

            # Расчет платежей
            monthly_rate = percent_clean / 100 / 12
            annuity_coeff = (monthly_rate * (1 + monthly_rate) ** term_clean) / ((1 + monthly_rate) ** term_clean - 1)
            monthly_payment = round(amount_clean * annuity_coeff, 2)
            initial_payment_percent = 20  # Стандартный начальный взнос 20%
            initial_payment = round(amount_clean * initial_payment_percent / 100, 2)
            total_payment = round(monthly_payment * term_clean, 2)
            total_overpayment = round(total_payment - amount_clean, 2)

            # Создаем папку для договоров
            contracts_dir = "credit_contracts"
            os.makedirs(contracts_dir, exist_ok=True)

            # Создаем документ
            doc = Document()

            # Стили форматирования
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Заголовок договора
            title = doc.add_paragraph()
            title_run = title.add_run(f'КРЕДИТНЫЙ ДОГОВОР № {loan_id}')
            title_run.bold = True
            title_run.font.size = Pt(14)
            title.alignment = 1  # Центральное выравнивание

            doc.add_paragraph(f"г. Москва, {datetime.now().strftime('%d.%m.%Y')}", style='Intense Quote')
            doc.add_paragraph()

            # Форматирование денежных сумм
            def format_money(value):
                return f"{value:,.2f}".replace(",", " ").replace(".", ",") + " руб."

            # 1. Стороны договора
            doc.add_heading('1. СТОРОНЫ ДОГОВОРА', level=1)
            doc.add_paragraph("1.1. КРЕДИТОР:", style='List Bullet')
            creditor = [
                "АО «Коммерческий Банк ФинансКредит»",
                "ИНН 7701234567, КПП 770101001",
                "ОГРН 1234500000123",
                "адрес: 123456, г. Москва, ул. Тверская, д. 1",
                "р/с 40702810100000001234 в АО «ФинансКредит»",
                "к/с 30101810000000000234",
                "БИК 044525234"
            ]
            for line in creditor:
                doc.add_paragraph(line, style='List Bullet 2')

            doc.add_paragraph("1.2. ЗАЕМЩИК:", style='List Bullet')
            debtor = [
                client_data[0],
                f"адрес: {client_data[1]}",
                f"телефон: {client_data[2]}",
                f"ИНН {client_data[3]}",
            ]
            for line in debtor:
                doc.add_paragraph(line, style='List Bullet 2')

            doc.add_paragraph()

            # 2. Условия кредитования
            doc.add_heading('2. УСЛОВИЯ КРЕДИТОВАНИЯ', level=1)
            terms = [
                f"2.1. Сумма кредита: {format_money(amount_clean)}",
                f"2.2. Процентная ставка: {percent_clean}% годовых",
                f"2.3. Срок кредита: {term_clean} месяцев",
                f"2.4. Начальный взнос: {format_money(initial_payment)} ({initial_payment_percent}% от суммы)",
                f"2.5. Ежемесячный платеж: {format_money(monthly_payment)}",
                f"2.6. Общая сумма выплат: {format_money(total_payment)}",
                f"2.7. Переплата за период: {format_money(total_overpayment)}"
            ]
            for term in terms:
                doc.add_paragraph(term)

            # 3. График платежей (первые 3 месяца)
            doc.add_paragraph("2.8. График платежей (фрагмент):")
            table = doc.add_table(rows=4, cols=4)
            table.style = 'Table Grid'

            # Заголовки таблицы
            headers = ['Дата платежа', 'Сумма платежа', 'Основной долг', 'Проценты']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].bold = True

            # Заполняем таблицу
            for month in range(1, 4):
                payment_date = (datetime.now() + timedelta(days=30 * month)).strftime('%d.%m.%Y')
                principal = round(amount_clean / term_clean, 2)
                interest = round(monthly_payment - principal, 2)

                row = table.rows[month]
                row.cells[0].text = payment_date
                row.cells[1].text = format_money(monthly_payment)
                row.cells[2].text = format_money(principal)
                row.cells[3].text = format_money(interest)

            doc.add_paragraph("Полный график платежей прилагается в Приложении 1")
            doc.add_page_break()

            # 4. Общие условия
            doc.add_heading('3. ОБЩИЕ УСЛОВИЯ', level=1)
            conditions = [
                "3.1. Кредит предоставляется на условиях полного возврата.",
                "3.2. Досрочное погашение возможно после 6 месяцев обслуживания кредита.",
                "3.3. При просрочке платежа начисляется пеня 0,1% от суммы просрочки за каждый день.",
                "3.4. При нарушении условий договора процентная ставка может быть увеличена до 20% годовых."
            ]
            for condition in conditions:
                doc.add_paragraph(condition, style='List Number')

            # 5. Подписи сторон
            doc.add_heading('4. ПОДПИСИ СТОРОН', level=1)
            doc.add_paragraph("КРЕДИТОР: ___________________/Петров А.И./")
            doc.add_paragraph("М.П.")
            doc.add_paragraph()
            doc.add_paragraph("ЗАЕМЩИК: ___________________/___________/")
            doc.add_paragraph("М.П.")

            # Сохраняем документ
            filename = f"{contracts_dir}/КД №{loan_id} - {client_data[0]}.docx"
            doc.save(filename)

            # Показываем путь к файлу
            abs_path = os.path.abspath(filename)
            messagebox.showinfo(
                "Договор создан",
                f"Кредитный договор успешно сформирован:\n\n{abs_path}"
            )

        except Exception as e:
            messagebox.showerror(
                "Ошибка генерации",
                f"Произошла ошибка при создании договора:\n\n{str(e)}"
            )

    def add_to_db(self):
        """Добавление кредита в базу данных"""
        en_id = self.en_id_entry.get()
        amount = self.amount_entry.get()
        percent = self.percent_entry.get()
        term = self.term_entry.get()
        start_date = datetime.now().strftime("%Y-%m-%d")

        try:
            # Подключение к базе данных
            conn = mysql.connector.connect(
                host='localhost',
                user='root',
                password='22345267',
                database='loan_db'
            )
            cursor = conn.cursor()

            # Проверка существования клиента
            cursor.execute("SELECT COUNT(*) FROM legal_entities WHERE entitie_id = %s", (en_id,))
            client_exists = cursor.fetchone()[0]

            if client_exists == 0:
                messagebox.showerror("Ошибка", "Клиент с указанным ID не существует.")
                conn.close()
                return

            # Добавление кредита в базу данных
            cursor.execute(
                "INSERT INTO loans (entitie_id, amount, percent, start_date,  term) "
                "VALUES (%s, %s, %s, %s, %s)",
                (en_id, amount, percent, start_date, term))

            # Получаем ID созданного кредита
            loan_id = cursor.lastrowid
            conn.commit()

            # Получаем данные клиента для договора
            cursor.execute("""
                SELECT en_name, address, phone_number, inn 
                FROM legal_entities 
                WHERE entitie_id = %s
            """, (en_id,))
            client_data = cursor.fetchone()

            messagebox.showinfo("Успех", "Кредит успешно добавлен в базу данных.")

            # Предложение экспортировать договор
            if messagebox.askyesno("Экспорт договора", "Экспортировать кредитный договор?"):
                self.generate_credit_contract(loan_id, en_id, client_data, amount, percent, term)

        except mysql.connector.Error as err:
            messagebox.showerror("Ошибка базы данных", f"Ошибка: {err}")
        finally:
            if 'conn' in locals() and conn.is_connected():
                conn.close()

class Show():
    def __init__(self):
        def export_payment():
            """Экспорт полного графика платежей для выбранного кредита"""
            selected_items = table2.selection()

            if not selected_items:
                messagebox.showwarning("Предупреждение", "Выберите кредит из таблицы")
                return

            selected_item = table2.selection()[0]
            loan_data = table2.item(selected_item)["values"]

            if not loan_data:
                messagebox.showerror("Ошибка", "Не удалось получить данные кредита")
                return

            loan_id = loan_data[0]
            amount = float(loan_data[2])
            percent = float(loan_data[3])
            term = int(loan_data[5])

            try:
                # Получаем данные клиента
                conn = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='22345267',
                    database='loan_db'
                )
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT en_name FROM legal_entities 
                    WHERE entitie_id = %s
                """, (loan_data[1],))
                client_name = cursor.fetchone()[0]
                conn.close()

                # Рассчитываем платежи
                monthly_rate = percent / 100 / 12
                annuity_coeff = (monthly_rate * (1 + monthly_rate) ** term) / ((1 + monthly_rate) ** term - 1)
                monthly_payment = round(amount * annuity_coeff, 2)

                # Создаем документ
                doc = Document()

                # Стили форматирования
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)

                # Заголовок
                title = doc.add_paragraph()
                title_run = title.add_run(f'Приложение №1 к Кредитному договору №{loan_id}')
                title_run.bold = True
                title_run.font.size = Pt(14)
                title.alignment = 1  # Центральное выравнивание

                doc.add_paragraph(f"График платежей по кредиту для {client_name}", style='Intense Quote')
                doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")
                doc.add_paragraph()

                # Форматирование денежных сумм
                def format_money(value):
                    return f"{value:,.2f}".replace(",", " ").replace(".", ",") + " руб."

                # Основные параметры кредита
                doc.add_paragraph(f"Сумма кредита: {format_money(amount)}")
                doc.add_paragraph(f"Процентная ставка: {percent}% годовых")
                doc.add_paragraph(f"Срок кредита: {term} месяцев")
                doc.add_paragraph(f"Ежемесячный платеж: {format_money(monthly_payment)}")
                doc.add_paragraph()

                # Создаем таблицу с графиком платежей
                table = doc.add_table(rows=term + 1, cols=5)
                table.style = 'Table Grid'

                # Заголовки таблицы
                headers = [
                    '№',
                    'Дата платежа',
                    'Сумма платежа',
                    'Основной долг',
                    'Проценты'
                ]

                for i, header in enumerate(headers):
                    table.cell(0, i).text = header
                    table.cell(0, i).paragraphs[0].runs[0].bold = True

                # Заполняем таблицу данными
                remaining_principal = amount

                for month in range(1, term + 1):
                    payment_date = (datetime.now() + timedelta(days=30 * month)).strftime('%d.%m.%Y')

                    # Расчет процентов на остаток
                    interest = round(remaining_principal * monthly_rate, 2)
                    principal = round(monthly_payment - interest, 2)

                    # Корректировка последнего платежа
                    if month == term:
                        principal = round(remaining_principal, 2)
                        monthly_payment = round(principal + interest, 2)

                    remaining_principal -= principal

                    # Заполняем строку таблицы
                    row = table.rows[month]
                    row.cells[0].text = str(month)
                    row.cells[1].text = payment_date
                    row.cells[2].text = format_money(monthly_payment)
                    row.cells[3].text = format_money(principal)
                    row.cells[4].text = format_money(interest)

                # Сохраняем документ
                os.makedirs("payment_schedules", exist_ok=True)
                filename = f"payment_schedules/График_платежей_КД_{loan_id}_{client_name}.docx"
                doc.save(filename)

                messagebox.showinfo(
                    "График платежей",
                    f"График платежей сохранен в файл:\n{os.path.abspath(filename)}"
                )

            except Exception as e:
                messagebox.showerror(
                    "Ошибка",
                    f"Не удалось сформировать график платежей:\n{str(e)}"
                )
        def open_lo_add():
            # Получаем выбранное юрлицо из таблицы
            selected_entity = table1.selection()
            entity_id = ""

            if selected_entity:
                entity_data = table1.item(selected_entity[0])["values"]
                if entity_data and len(entity_data) > 0:
                    entity_id = entity_data[0]  # Берем ID из первого столбца

            # Создаем окно добавления кредита и передаем ID юрлица
            open = Add_lo(entity_id)

        def open_en_add():
            open = Add_en()

        def update_treeview():
            table1.delete(*table1.get_children())
            table2.delete(*table2.get_children())

            # Установка соединения с базой данных MySQL
            conn = mysql.connector.connect(
                host='localhost',
                user='root',
                password='22345267',
                database='loan_db'
            )

            # Создание объекта cursor для выполнения SQL-запросов
            cursor1 = conn.cursor()

            # Выполнение SQL-запроса для выборки записей из таблицы
            cursor1.execute("SELECT * FROM legal_entities")

            # Получение всех выбранных записей
            en_lst_upd = cursor1.fetchall()

            # Создание объекта cursor для выполнения SQL-запросов
            cursor2 = conn.cursor()

            # Выполнение SQL-запроса для выборки записей из таблицы
            cursor2.execute("""
            SELECT l.loan_id, le.en_name, l.amount, l.percent, l.start_date, l.term 
            FROM loans l
            JOIN legal_entities le ON l.entitie_id = le.entitie_id
            """)

            # Получение всех выбранных записей
            lo_lst_upd = cursor2.fetchall()

            # Закрытие соединения с базой данных
            conn.close()

            for row in lo_lst_upd:
                table2.insert('', tkinter.END, values=row)

            for row in en_lst_upd:
                table1.insert('', tkinter.END, values=row)

        def search_en():
            search_value = en_searchfield.get()

            if search_value == '':
                messagebox.showerror(title='Ошибка', message='Поле должнобыть заполнено для поиска.')

            else:
                # Установка соединения с базой данных MySQL
                conn = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='22345267',
                    database='loan_db'
                )

                # Создание объекта cursor для выполнения SQL-запросов
                cursor = conn.cursor()

                query = "SELECT * FROM legal_entities WHERE entitie_id LIKE %s OR en_name LIKE %s OR address LIKE %s"
                cursor.execute(query, (
                    '%' + search_value + '%', '%' + search_value + '%', '%' + search_value + '%'))

                # Получение всех выбранных записей
                serched_en = cursor.fetchall()
                if serched_en == []:
                    messagebox.showinfo(title='Результат', message='Записи с такими данными не найдено.')

                else:
                    # Закрытие соединения с базой данных
                    conn.close()

                    table1.delete(*table1.get_children())

                    for row in serched_en:
                        table1.insert('', tkinter.END, values=row)

        def search_lo():
            search_value = lo_searchfield.get()

            if search_value == '':
                messagebox.showerror(title='Ошибка', message='Поле должнобыть заполнено для поиска.')
            else:
                # Установка соединения с базой данных MySQL
                conn = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='22345267',
                    database='loan_db'
                )

                # Создание объекта cursor для выполнения SQL-запросов
                cursor = conn.cursor()

                query = "SELECT * FROM loans WHERE loan_id LIKE %s OR entitie_id LIKE %s"
                cursor.execute(query, (
                '%' + search_value + '%', '%' + search_value + '%'))

                # Получение всех выбранных записей
                serched_lo = cursor.fetchall()

                # Закрытие соединения с базой данных
                conn.close()

                if serched_lo == []:
                    messagebox.showinfo(title='Результат', message='Записи с такими ID не найдено.')
                else:
                    table2.delete(*table2.get_children())

                    for row in serched_lo:
                        table2.insert('', tkinter.END, values=row)

        def delete_en():
            selected_items = table1.selection()

            if not selected_items:
                messagebox.showwarning("Предупреждение", "Пожалуйста, выберите элемент для удаления.")
                return  # Прерываем выполнение функции, если элемент не выбран

            selected_item = table1.selection()[0]
            item_values = table1.item(selected_item)["values"]
            en_id = item_values[0]

            responce = messagebox.askyesno(title='Подтверждение', message='Вы действительно хотите удалить запись?')
            if responce:
                conn = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='22345267',
                    database='loan_db'
                )

                cursor = conn.cursor()

                # Выполнение SQL-запроса для удаления записи из таблицы
                command = "DELETE FROM legal_entities WHERE entitie_id = %s"
                cursor.execute(command, (en_id,))

                # Применяем изменения
                conn.commit()

                # Закрытие соединения с базой данных
                conn.close()

                messagebox.showinfo("Удаление", "Запись успешно удалена.")

                update_treeview()
            else:
                # Код, который нужно выполнить, если пользователь отказался от удаления
                messagebox.showinfo("Удаление", "Удаление отменено.")

        def delete_lo():
            selected_items = table2.selection()

            if not selected_items:
                messagebox.showwarning("Предупреждение", "Пожалуйста, выберите элемент для удаления.")
                return  # Прерываем выполнение функции, если элемент не выбран

            selected_item = table2.selection()[0]
            item_values = table2.item(selected_item)["values"]
            lo_id = item_values[0]

            responce = messagebox.askyesno(title='Подтверждение', message='Вы действительно хотите удалить запись?')
            if responce:
                conn = mysql.connector.connect(
                    host='localhost',
                    user='root',
                    password='22345267',
                    database='loan_db'
                )

                cursor = conn.cursor()

                # Выполнение SQL-запроса для удаления записи из таблицы
                command = "DELETE FROM loans WHERE loan_id = %s"
                cursor.execute(command, (lo_id,))

                # Применяем изменения
                conn.commit()

                # Закрытие соединения с базой данных
                conn.close()

                messagebox.showinfo("Удаление", "Запись успешно удалена.")

                update_treeview()
            else:
                # Код, который нужно выполнить, если пользователь отказался от удаления
                messagebox.showinfo("Удаление", "Удаление отменено.")

        def update_en():
            selected_items = table1.selection()

            if not selected_items:
                messagebox.showwarning("Предупреждение", "Пожалуйста, выберите элемент для изменеия.")
                return  # Прерываем выполнение функции, если элемент не выбран

            selected_item = table1.selection()[0]

            item_values = table1.item(selected_item)["values"]
            en_id = item_values[0]
            en_name = item_values[1]
            en_address = item_values[2]
            en_phone = item_values[3]
            en_inn = item_values[4]

            def update():
                    responce = messagebox.askyesno(title='Подтверждение', message='Вы действительно хотите изменить запись?')
                    if responce:
                        upd_name = name_entry.get()
                        upd_address = address_entry.get()
                        upd_phone = phone_entry.get()

                        # Установка соединения с базой данных MySQL
                        conn = mysql.connector.connect(
                            host='localhost',
                            user='root',
                            password='22345267',
                            database='loan_db'
                        )

                        # Создание курсора для выполнения SQL-запросов
                        cursor1 = conn.cursor()

                        # Выполнение SQL-запроса для удаления записи из таблицы
                        command = "UPDATE legal_entities SET en_name = %s, address = %s, phone_number = %s WHERE entitie_id = %s"
                        cursor1.execute(command, (upd_name, upd_address, upd_phone,  en_id,))

                        # Применение изменений
                        conn.commit()

                        conn.close()

                        messagebox.showinfo("Изменение.", "Запись успешно обновлена.")

                        update_treeview()

                        entitie_upd.destroy()

            def format_phone(event=None):
                # Форматирует номер по маске ##-##-##
                current = phone_entry.get().replace("-", "")
                formatted = ""

                if len(current) > 0:
                    formatted = current[:2]  # Первые 2 цифры
                if len(current) > 2:
                    formatted += "-" + current[2:4]  # Добавляем - и следующие 2 цифры
                if len(current) > 4:
                    formatted += "-" + current[4:6]  # Добавляем - и последние 2 цифры

                # Обновляем поле ввода
                phone_entry.delete(0, "end")
                phone_entry.insert(0, formatted)

            def validate_inn_input(new_value):
                "Проверка ввода ИНН (только цифры, максимум 12 символов)"
                return new_value.isdigit() and len(new_value) <= 12 or new_value == ""

            entitie_upd = Tk()
            entitie_upd.title("Изменение юр. лица.")
            entitie_upd['bg'] = '#d7cecc'
            entitie_upd.geometry('350x300')
            entitie_upd.resizable(width=False, height=False)

            name_label = Label(entitie_upd, text="Название", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            name_label.pack()
            name_entry = Entry(entitie_upd, bg='#f0f6f6')
            name_entry.pack()
            name_entry.insert(0, en_name)

            address_label = Label(entitie_upd, text="Адрес", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            address_label.pack()
            address_entry = Entry(entitie_upd, bg='#f0f6f6')
            address_entry.pack()
            address_entry.insert(0, en_address)

            phone_label = Label(entitie_upd, text="Номер телефона", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            phone_label.pack()
            phone_entry = Entry(entitie_upd, bg='#f0f6f6')
            phone_entry.bind("<KeyRelease>", format_phone)
            phone_entry.pack()
            phone_entry.insert(0, en_phone)

            inn_label = Label(entitie_upd, text="ИНН", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            inn_label.pack()
            inn_entry = Entry(entitie_upd, bg='#f0f6f6', validate="key")
            inn_entry['validatecommand'] = (entitie_upd.register(validate_inn_input), '%P')
            inn_entry.pack()
            inn_entry.insert(0, en_inn)

            add_button = Button(entitie_upd, text="Обновить", command=update, bg='#f0f6f6')
            add_button.pack(padx=10, pady=8)

        def update_lo():
            selected_items = table2.selection()

            if not selected_items:
                messagebox.showwarning("Предупреждение", "Пожалуйста, выберите элемент для изменеия.")
                return  # Прерываем выполнение функции, если элемент не выбран

            selected_item = table2.selection()[0]

            item_values = table2.item(selected_item)["values"]
            lo_id = item_values[0]
            en_id = item_values[1]
            amount = item_values[2]
            percent = item_values[3]
            term = item_values[5]

            def update():
                    upd_en_id = en_id_entry.get()
                    upd_amount = amount_entry.get()
                    upd_percent = percent_entry.get()
                    upd_term = term_entry.get()

                    # Установка соединения с базой данных MySQL
                    conn = mysql.connector.connect(
                        host='localhost',
                        user='root',
                        password='22345267',
                        database='loan_db'
                    )

                    # Создание курсора для выполнения SQL-запросов
                    cursor1 = conn.cursor()

                    # Операция проверк исуществования клиента
                    cursor1.execute("SELECT COUNT(*) FROM legal_entities WHERE entitie_id = %s", (upd_en_id,))
                    client_exists = cursor1.fetchone()[0]

                    if client_exists == 0:
                        messagebox.showerror("Ошибка", "Клиент с указанным ID не существует.")

                    else:
                        responce = messagebox.askyesno(title='Подтверждение', message='Вы действительно хотите изменить запись?')
                        if responce:
                            # Выполнение SQL-запроса для удаления записи из таблицы
                            command = ("UPDATE loans SET entitie_id = %s, amount = %s, percent = %s, term = %s WHERE loan_id = %s")
                            cursor1.execute(command, (upd_en_id, upd_amount, upd_percent, upd_term, lo_id))

                            messagebox.showinfo("Изменение.", "Запись успешно обновлена.")

                            loan_upd.destroy()

                        # Применение изменений
                        conn.commit()

                        update_treeview()

                        # Закрытие соединения с базой данных
                        conn.close()


            loan_upd = Tk()
            loan_upd.title("Изменение кредита.")
            loan_upd['bg'] = '#d7cecc'
            loan_upd.geometry('350x300')
            loan_upd.resizable(width=False, height=False)

            en_id_label = Label(loan_upd, text="ID лица", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            en_id_label.pack()
            en_id_entry = Entry(loan_upd, bg='#f0f6f6')
            en_id_entry.pack()
            en_id_entry.insert(0, en_id)

            amount_label = Label(loan_upd, text="Сумма", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            amount_label.pack()
            amount_entry = Entry(loan_upd, bg='#f0f6f6')
            amount_entry.pack()
            amount_entry.insert(0, amount)

            percent_label = Label(loan_upd, text="Процент", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            percent_label.pack()
            percent_entry = Entry(loan_upd, bg='#f0f6f6')
            percent_entry.pack()
            percent_entry.insert(0, percent)

            term_label = Label(loan_upd, text="Срок", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
            term_label.pack()
            term_entry = Entry(loan_upd, bg='#f0f6f6')
            term_entry.pack()
            term_entry.insert(0, term)

            add_button = Button(loan_upd, text="Обновить", command=update, bg='#f0f6f6')
            add_button.pack(padx=10, pady=8)

        show = Tk()
        show.title("Просмотр записей.")
        show['bg'] = '#d7cecc'
        show.geometry('800x350')
        show.resizable(width=False, height=False)

        # Главный контейнер с Notebook
        tab_control = ttk.Notebook(show)
        tab_control.place(relwidth=1.0, relheight=1.0)

        # Вкладка "Лица"
        tab1 = Frame(tab_control, bg='#d7cecc')
        tab_control.add(tab1, text='Лица')

        # Вкладка "Кредиты"
        tab2 = Frame(tab_control, bg='#d7cecc')
        tab_control.add(tab2, text='Кредиты')

        # Подключение к БД
        conn = mysql.connector.connect(
            host='localhost',
            user='root',
            password='22345267',
            database='loan_db'
        )
        cursor = conn.cursor()

        # Заголовок
        lb1 = Label(tab1, text='Лица', font='Arial 12 bold', bg='#d7cecc')
        lb1.place(relx=0.02, rely=0.02)

        # Таблица
        cursor.execute("SELECT * FROM legal_entities")
        en_lst = cursor.fetchall()

        heads = ['id', 'Название', 'Адрес', 'Телефонный номер', 'ИНН']
        table1 = ttk.Treeview(tab1, show='headings', height=10)
        table1['columns'] = heads
        for row in en_lst:
            table1.insert('', tkinter.END, values=row)
        for header in heads:
            table1.heading(header, text=header, anchor='center')
            table1.column(header, anchor='center')
        table1.column('id', width=50)
        table1.column('Название', width=120)
        table1.column('Адрес', width=200)
        table1.column('Телефонный номер', width=120)
        table1.column('ИНН', width=120)
        table1.place(relx=0.02, rely=0.1, relwidth=0.96)

        control_frame1 = Frame(tab1, bg='#d7cecc')
        control_frame1.place(relx=0.02, rely=0.8, relwidth=0.96)

        # Кнопки управления
        buttons_frame1 = Frame(control_frame1, bg='#d7cecc')
        buttons_frame1.pack(side=LEFT)

        open_en_add = Button(buttons_frame1, text='Добавить', command=open_en_add, bg='#f0f6f6')
        open_en_add.pack(side=LEFT, padx=5, pady=20)
        en_change = Button(buttons_frame1, text='Изменить', command=update_en, bg='#f0f6f6')
        en_change.pack(side=LEFT, padx=5, pady=20)
        en_del = Button(buttons_frame1, text='Удалить', command=delete_en, bg='#f0f6f6')
        en_del.pack(side=LEFT, padx=5, pady=20)

        # Блок поиска
        search_frame1 = Frame(control_frame1, bg='#d7cecc')
        search_frame1.pack(side=RIGHT)

        en_update = Button(search_frame1, text='Обновить', command=update_treeview, bg='#f0f6f6')
        en_update.pack(side=RIGHT, padx=5, pady=20)
        en_searchbutton = Button(search_frame1, text='Поиск', command=search_en, bg='#f0f6f6')
        en_searchbutton.pack(side=RIGHT, padx=5, pady=20)
        en_searchfield = Entry(search_frame1, width=40, bg='#f0f6f6')
        en_searchfield.pack(side=RIGHT, padx=5, pady=20)

        # Заголовок
        lb2 = Label(tab2, text='Кредиты', font='Arial 12 bold', bg='#d7cecc')
        lb2.place(relx=0.02, rely=0.02)

        # Таблица
        cursor.execute("""
            SELECT l.loan_id, le.en_name, l.amount, l.percent, l.start_date, l.term 
            FROM loans l
            JOIN legal_entities le ON l.entitie_id = le.entitie_id
        """)
        lo_lst = cursor.fetchall()

        heads = ['id кредита', 'Заемщик', 'Сумма', 'Процент', 'Дата открытия', 'Срок']
        table2 = ttk.Treeview(tab2, show='headings', height=10)
        table2['columns'] = heads
        for row in lo_lst:
            table2.insert('', tkinter.END, values=row)
        for header in heads:
            table2.heading(header, text=header, anchor='center')
            table2.column(header, anchor='center')
        table2.column('id кредита', width=70)
        table2.column('Заемщик', width=150)  # Шире для названий
        table2.column('Сумма', width=150)
        table2.column('Процент', width=70)
        table2.column('Дата открытия', width=100)
        table2.column('Срок', width=120)
        table2.place(relx=0.02, rely=0.1, relwidth=0.96)

        # Контейнер для кнопок управления и поиска
        control_frame2 = Frame(tab2, bg='#d7cecc')
        control_frame2.place(relx=0.02, rely=0.8, relwidth=0.96)

        # Кнопки управления
        buttons_frame2 = Frame(control_frame2, bg='#d7cecc')
        buttons_frame2.pack(side=LEFT)

        open_lo_add = Button(buttons_frame2, text='Добавить', command=open_lo_add, bg='#f0f6f6')
        open_lo_add.pack(side=LEFT, padx=5, pady=20)
        lo_change = Button(buttons_frame2, text='Изменить', command=update_lo, bg='#f0f6f6')
        lo_change.pack(side=LEFT, padx=5, pady=20)
        lo_del = Button(buttons_frame2, text='Удалить', command=delete_lo, bg='#f0f6f6')
        lo_del.pack(side=LEFT, padx=5, pady=20)
        lo_export = Button(buttons_frame2, text='Рассчитать', command=export_payment, bg='#f0f6f6')
        lo_export.pack(side=LEFT, padx=5, pady=20)

        # Блок поиска
        search_frame2 = Frame(control_frame2, bg='#d7cecc')
        search_frame2.pack(side=RIGHT)

        lo_update = Button(search_frame2, text='Обновить', command=update_treeview, bg='#f0f6f6')
        lo_update.pack(side=RIGHT, padx=5, pady=20)
        lo_searchbutton = Button(search_frame2, text='Поиск', command=search_lo, bg='#f0f6f6')
        lo_searchbutton.pack(side=RIGHT, padx=5, pady=20)
        lo_searchfield = Entry(search_frame2, width=40, bg='#f0f6f6')
        lo_searchfield.pack(side=RIGHT, padx=5, pady=20)

        # Закрытие соединения
        conn.close()

def login():
    username = login_entry.get()
    password = password_entry.get()

    conn = mysql.connector.connect(
        host='localhost',
        user='root',
        password='22345267',
        database='loan_db'
    )

    cursor = conn.cursor(dictionary=True)
    query = "SELECT * FROM employees WHERE login = %s AND password = %s"
    cursor.execute(query, (username, password))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

    if user:
        # Успешная авторизация
        messagebox.showinfo("Успех", "Авторизация успешна!")
        show = Show()
        root.destroy()
    else:
        # Неверный логин или пароль
        messagebox.showerror("Ошибка", "Неверный логин или пароль.")

root = Tk()
root.title("Авторизация.")
root['bg'] = '#d7cecc'
root.geometry('450x200')
root.resizable(width=False, height=False)

main_label = Label(root, text='Авторизация', font='Arial 15 bold', bg='#d7cecc')
main_label.pack()

login_label = Label(root, text="Логин", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
login_label.pack()

login_entry = Entry(root, bg='#fafafa', font='arial 12')
login_entry.pack()

password_label = Label(root, text="Пароль", font='Arial 11 bold', bg='#d7cecc', padx=10, pady=8)
password_label.pack()

password_entry = Entry(root, show="*", bg='#fafafa', font='arial 12')
password_entry.pack()

login_button = Button(root, text="Войти", command=login, bg='#f0f6f6')
login_button.pack(padx=10, pady=8)

root.mainloop()
